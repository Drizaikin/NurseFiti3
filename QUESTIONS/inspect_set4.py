import subprocess, sys
try:
    from docx import Document
except ImportError:
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
    from docx import Document

doc = Document(r'C:\Users\User\Desktop\NurseFiti\NurseFiti3\QUESTIONS\NCK_MCQ_Compilation_Set4_Answers_Rationales.docx')

print("=== TABLES in document ===")
print(f"Number of tables: {len(doc.tables)}")
for i, table in enumerate(doc.tables[:3]):
    print(f"\nTable {i}: {len(table.rows)} rows x {len(table.columns)} cols")
    for row in table.rows[:5]:
        print("  Row:", [cell.text[:60] for cell in row.cells])

print("\n=== PARAGRAPHS (first 80, showing all properties) ===")
for i, para in enumerate(doc.paragraphs[:80]):
    text = para.text.strip()
    if not text:
        continue
    is_bold = bool(para.runs) and all(r.bold for r in para.runs if r.text.strip())
    is_red = False
    for r in para.runs:
        if r.text.strip():
            try:
                rgb = r.font.color.rgb
                if rgb:
                    is_red = True
                    break
            except:
                pass
    style = para.style.name if para.style else "None"
    print(f"[{i:3d}] bold={is_bold} red={is_red} style={style!r:30s} | {text[:90]}")
